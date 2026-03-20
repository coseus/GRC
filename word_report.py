from io import BytesIO

from docx import Document

from utils import sort_recommendations_by_risk


def generate_word_report(
    company_name,
    assessment_name,
    assessment_date,
    overall_score,
    domain_scores,
    summary,
    recommendations,
    mapping_rows,
    responses=None,
    include_proof=False,
    include_mapping=False,
    export_mode="Executive",
    management_content=None,
):
    recommendations = sort_recommendations_by_risk(recommendations)
    responses = responses or []
    management_content = management_content or {
        "intro": "",
        "top_actions": [],
        "domain_blocks": {},
    }

    doc = Document()

    doc.add_heading("Cyber Security Assessment Report", 0)
    doc.add_paragraph(f"Company: {company_name}")
    doc.add_paragraph(f"Assessment: {assessment_name}")
    doc.add_paragraph(f"Date: {assessment_date}")
    doc.add_paragraph(f"Export mode: {export_mode}")
    doc.add_paragraph(f"Overall Score: {overall_score:.1f}/100")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary or "-")

    doc.add_heading("Domain Scores", level=1)
    if domain_scores:
        score_table = doc.add_table(rows=1, cols=3)
        score_table.style = "Table Grid"
        hdr = score_table.rows[0].cells
        hdr[0].text = "Domain"
        hdr[1].text = "Score"
        hdr[2].text = "Maturity"

        for domain, score in domain_scores.items():
            row = score_table.add_row().cells
            row[0].text = domain
            row[1].text = f"{score:.1f}"
            row[2].text = (
                "Optimized" if score >= 90 else
                "Managed" if score >= 75 else
                "Defined" if score >= 55 else
                "Repeatable" if score >= 30 else
                "Initial"
            )
    else:
        doc.add_paragraph("-")

    if export_mode == "Executive":
        intro = (management_content.get("intro") or "").strip()
        top_actions = management_content.get("top_actions", [])
        domain_blocks = management_content.get("domain_blocks", {})

        if intro:
            doc.add_paragraph(intro)

        doc.add_heading("Top 5 Priority Actions", level=1)
        if top_actions:
            for item in top_actions[:5]:
                if str(item).strip():
                    doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph("-")

        doc.add_heading("Key Areas of Improvement", level=1)
        if domain_blocks:
            for domain, bullets in domain_blocks.items():
                doc.add_paragraph(domain, style="Heading 2")
                for bullet in bullets:
                    if str(bullet).strip():
                        doc.add_paragraph(bullet, style="List Bullet")
        else:
            doc.add_paragraph("-")

    else:
        doc.add_heading("Recommendations", level=1)
        if recommendations:
            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Domain"
            hdr[1].text = "Risk"
            hdr[2].text = "Recommendation"
            hdr[3].text = "Status"
            hdr[4].text = "Responsible"
            hdr[5].text = "Deadline"

            for r in recommendations:
                row = table.add_row().cells
                row[0].text = r.get("domain_name", "")
                row[1].text = r.get("risk", "")
                row[2].text = r.get("text", "")
                row[3].text = r.get("status", "")
                row[4].text = r.get("responsible", "") or ""
                row[5].text = r.get("deadline", "") or ""
        else:
            doc.add_paragraph("-")

    if export_mode == "Detailed":
        doc.add_heading("Assessment Details", level=1)
        if responses:
            cols = 6 if include_proof else 5
            details_table = doc.add_table(rows=1, cols=cols)
            details_table.style = "Table Grid"
            hdr = details_table.rows[0].cells
            hdr[0].text = "Domain"
            hdr[1].text = "Question"
            hdr[2].text = "Answer"
            hdr[3].text = "Score"
            hdr[4].text = "Notes"
            if include_proof:
                hdr[5].text = "Proof"

            for r in responses:
                row = details_table.add_row().cells
                row[0].text = r.get("domain", "")
                row[1].text = r.get("question", "")
                row[2].text = str(r.get("answer_value", "") or "")
                row[3].text = str(r.get("score", "") or "")
                row[4].text = r.get("notes", "") or ""
                if include_proof:
                    row[5].text = r.get("proof", "") or ""
        else:
            doc.add_paragraph("-")

    if include_mapping:
        doc.add_heading("Control Mapping", level=1)
        if mapping_rows:
            map_table = doc.add_table(rows=1, cols=6)
            map_table.style = "Table Grid"
            hdr = map_table.rows[0].cells
            hdr[0].text = "Domain"
            hdr[1].text = "Question"
            hdr[2].text = "ISO 27001"
            hdr[3].text = "NIST CSF"
            hdr[4].text = "CIS"
            hdr[5].text = "NIS2"

            for r in mapping_rows:
                row = map_table.add_row().cells
                row[0].text = r.get("domain", "")
                row[1].text = r.get("question", "")
                row[2].text = r.get("iso27001", "")
                row[3].text = r.get("nist_csf", "")
                row[4].text = r.get("cis_control", "")
                row[5].text = r.get("nis2", "")
        else:
            doc.add_paragraph("-")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
